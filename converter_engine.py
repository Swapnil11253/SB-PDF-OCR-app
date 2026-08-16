import os
import glob
import re
import json
import copy
import io
import base64
import time
import shutil
import subprocess
import docx

try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn as docx_qn
from lxml import etree
from pptx import Presentation
from pptx.oxml.ns import qn as pptx_qn
from pptx.enum.text import MSO_AUTO_SIZE, PP_ALIGN
from pptx.util import Inches as PPTX_Inches

try:
    import latex2mathml.converter as _l2m
    import mathml2omml as _m2o
    NATIVE_ENGINE_AVAILABLE = True
except ImportError:
    NATIVE_ENGINE_AVAILABLE = False

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

_CONV_STATS = {"success": 0, "failed": 0, "failed_list": []}

def _reset_stats():
    _CONV_STATS["success"] = 0
    _CONV_STATS["failed"] = 0
    _CONV_STATS["failed_list"] = []

_ACCENTED_ONE_RE = re.compile(
    r'\\(hat|widehat|vec|overrightarrow|tilde|widetilde)'
    r'(?:\s*\{\s*1\s*\}|\s+1(?![0-9]))'
)

def _fix_ocr_unit_vector_ones(s):
    letters = ['i', 'j', 'k']
    counter = {'n': 0}

    def _sub(m):
        cmd = m.group(1)
        letter = letters[counter['n'] % len(letters)]
        counter['n'] += 1
        return f'\\{cmd}{{{letter}}}'

    return _ACCENTED_ONE_RE.sub(_sub, s)

def sanitize_latex(latex_str):
    s = latex_str.strip()
    s = re.sub(r'\{\\rm\s+([^{}]+)\}', r'\\mathrm{\1}', s)
    s = re.sub(r'\{\\bf\s+([^{}]+)\}', r'\\mathbf{\1}', s)
    s = re.sub(r'\\hbox\s*\{([^{}]*)\}', r'\\text{\1}', s)
    s = s.replace(r'\{cdot\}', r'\cdot').replace(r'\;', ' ').replace(r'\,', ' ')
    s = re.sub(
        r'([A-Za-z])(?:_\{?([A-Za-z0-9]+)\}?)?[\u2192\u20d7]',
        lambda m: (r'\vec{' + m.group(1) + '_{' + m.group(2) + '}}') if m.group(2)
                  else (r'\vec{' + m.group(1) + '}'),
        s
    )
    s = _fix_ocr_unit_vector_ones(s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def aggressive_sanitize_latex(s):
    s = re.sub(r'\\d?frac', r'\\frac', s)
    s = re.sub(r'\\varDelta', r'\\Delta', s)
    s = re.sub(r'\\varGamma', r'\\Gamma', s)
    s = re.sub(r'\\bold\{([^{}]*)\}', r'\\mathbf{\1}', s)
    s = re.sub(r'\\text\s*\{\s*\}', '', s)
    s = s.replace(r'\displaystyle', '').replace(r'\textstyle', '')
    s = s.replace(r'\nonumber', '').replace(r'\notag', '')
    s = re.sub(r'\\label\{[^{}]*\}', '', s)
    s = re.sub(r'\\tag\{[^{}]*\}', '', s)
    s = s.replace(r'\;', ' ').replace(r'\!', '').replace(r'\quad', '  ').replace(r'\qquad', '    ')
    if s.count('{') != s.count('}'):
        if s.count('{') > s.count('}'):
            s = s.rstrip('{')
        else:
            s = s.lstrip('}')
    return re.sub(r'\s+', ' ', s).strip()

def _repair_omml(omml_str):
    omml_str = re.sub(
        r'(<m:pos[^/]*/>)</m:groupChr>(<m:e>)',
        r'\1</m:groupChrPr>\2',
        omml_str
    )
    return omml_str

_ARROW_TO_ACCENT_CHAR = {
    '\u2192': '\u20d7',
    '\u2190': '\u20d6',
    '\u2194': '\u20e1',
    '\u00af': '\u0305',
    '\u02c7': '\u030c',
}

_LIMUPP_LIM_CHAR_TO_ACCENT = {
    '^': '\u0302',
    '~': '\u0303',
    '\u02d9': '\u0307',
    '\u00a8': '\u0308',
    '\u02d8': '\u0306',
}

def _vector_groupchr_to_accent(omath_el):
    for grp in omath_el.findall(f'.//{{{MATH_NS}}}groupChr'):
        pr = grp.find(f'{{{MATH_NS}}}groupChrPr')
        if pr is None:
            continue
        chr_el = pr.find(f'{{{MATH_NS}}}chr')
        pos_el = pr.find(f'{{{MATH_NS}}}pos')
        if chr_el is None or pos_el is None:
            continue
        chr_val = chr_el.get(f'{{{MATH_NS}}}val')
        pos_val = pos_el.get(f'{{{MATH_NS}}}val')
        if pos_val != 'top' or chr_val not in _ARROW_TO_ACCENT_CHAR:
            continue
        e_el = grp.find(f'{{{MATH_NS}}}e')
        if e_el is None:
            continue
        parent = grp.getparent()
        idx = list(parent).index(grp)
        acc = etree.Element(f'{{{MATH_NS}}}acc')
        acc_pr = etree.SubElement(acc, f'{{{MATH_NS}}}accPr')
        new_chr = etree.SubElement(acc_pr, f'{{{MATH_NS}}}chr')
        new_chr.set(f'{{{MATH_NS}}}val', _ARROW_TO_ACCENT_CHAR[chr_val])
        acc.append(e_el)
        parent.remove(grp)
        parent.insert(idx, acc)
    return omath_el

def _limupp_accent_to_acc(omath_el):
    for limupp in omath_el.findall(f'.//{{{MATH_NS}}}limUpp'):
        lim_el = limupp.find(f'{{{MATH_NS}}}lim')
        e_el = limupp.find(f'{{{MATH_NS}}}e')
        if lim_el is None or e_el is None:
            continue
        t_el = lim_el.find(f'.//{{{MATH_NS}}}t')
        if t_el is None or t_el.text not in _LIMUPP_LIM_CHAR_TO_ACCENT:
            continue
        accent_char = _LIMUPP_LIM_CHAR_TO_ACCENT[t_el.text]
        parent = limupp.getparent()
        if parent is None:
            continue
        idx = list(parent).index(limupp)
        acc = etree.Element(f'{{{MATH_NS}}}acc')
        acc_pr = etree.SubElement(acc, f'{{{MATH_NS}}}accPr')
        new_chr = etree.SubElement(acc_pr, f'{{{MATH_NS}}}chr')
        new_chr.set(f'{{{MATH_NS}}}val', accent_char)
        acc.append(e_el)
        parent.remove(limupp)
        parent.insert(idx, acc)
    return omath_el

def _fix_radical_missing_degree(omath_el):
    for rad in omath_el.findall(f'.//{{{MATH_NS}}}rad'):
        if rad.find(f'{{{MATH_NS}}}deg') is not None:
            continue
        rad_pr = etree.Element(f'{{{MATH_NS}}}radPr')
        deg_hide = etree.SubElement(rad_pr, f'{{{MATH_NS}}}degHide')
        deg_hide.set(f'{{{MATH_NS}}}val', '1')
        empty_deg = etree.Element(f'{{{MATH_NS}}}deg')
        rad.insert(0, empty_deg)
        rad.insert(0, rad_pr)
    return omath_el

def _convert_via_native_engine(cleaned_latex):
    if not NATIVE_ENGINE_AVAILABLE:
        return None
    mathml = _l2m.convert(cleaned_latex)
    omml_str = _repair_omml(_m2o.convert(mathml))
    wrapped = f'<root xmlns:m="{MATH_NS}">{omml_str}</root>'
    root = etree.fromstring(wrapped.encode("utf-8"))
    omath_el = root.find(f'.//{{{MATH_NS}}}oMath')
    if omath_el is not None:
        omath_el = _vector_groupchr_to_accent(omath_el)
        omath_el = _limupp_accent_to_acc(omath_el)
        omath_el = _fix_radical_missing_degree(omath_el)
    return omath_el

def convert_latex_to_omml_native(latex_code):
    cleaned = sanitize_latex(latex_code)
    if not cleaned:
        return None
    try:
        el = _convert_via_native_engine(cleaned)
        if el is not None:
            _CONV_STATS["success"] += 1
            return el
    except Exception:
        pass

    try:
        retry_latex = aggressive_sanitize_latex(cleaned)
        if retry_latex and retry_latex != cleaned:
            el = _convert_via_native_engine(retry_latex)
            if el is not None:
                _CONV_STATS["success"] += 1
                return el
    except Exception:
        pass

    _CONV_STATS["failed"] += 1
    if len(_CONV_STATS["failed_list"]) < 15:
        _CONV_STATS["failed_list"].append(cleaned[:60])
    return None

def _new_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(11)
    return doc

def build_docx_from_text(input_text, output_docx_path):
    _reset_stats()
    doc = _new_doc()
    lines = input_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        # Check if line contains Math
        if line.startswith('$') and line.endswith('$'):
            clean_math = line.strip('$')
            omml_el = convert_latex_to_omml_native(clean_math)
            if omml_el is not None:
                p._p.append(omml_el)
            else:
                run = p.add_run(clean_math)
                run.font.italic = True
        else:
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            
    doc.save(output_docx_path)
    return _CONV_STATS
